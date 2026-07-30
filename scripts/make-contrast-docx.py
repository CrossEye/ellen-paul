"""Regenerate "Ellen vs Weir - Contrast Menu.docx" from content/weir-contrast-draft.md.

Usage: python scripts/make-contrast-docx.py

Tailored to the structures that file actually uses: #/##/### headings,
bullet lists, blockquotes (including a heading inside a blockquote and the
draft-copy quotes), horizontal rules, and **bold** / *italic* inline runs.
"""

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "content" / "weir-contrast-draft.md"
OUT = ROOT / "Ellen vs Weir - Contrast Menu.docx"

ACCENT = "B8860B"  # gold, to match the site's palette
INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def add_runs(par, text):
    """Split markdown inline emphasis into styled runs."""
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            par.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            par.add_run(token[1:-1]).italic = True
        else:
            par.add_run(token)


def set_left_border(par, color=ACCENT, size="18"):
    pbdr = par._p.get_or_add_pPr().makeelement(qn("w:pBdr"), {})
    left = pbdr.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): size,
        qn("w:space"): "10",
        qn("w:color"): color,
    })
    pbdr.append(left)
    par._p.get_or_add_pPr().append(pbdr)


def add_rule(doc):
    par = doc.add_paragraph()
    pbdr = par._p.get_or_add_pPr().makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "999999",
    })
    pbdr.append(bottom)
    par._p.get_or_add_pPr().append(pbdr)


def add_quote(doc, text, heading=False):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.3)
    par.paragraph_format.space_after = Pt(6)
    set_left_border(par)
    if heading:
        run = par.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x8B, 0x66, 0x08)
    else:
        add_runs(par, text)


def build():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "---":
            add_rule(doc)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], 0)
        elif stripped.startswith("## "):
            add_runs(doc.add_heading("", 1), stripped[3:])
        elif stripped.startswith("### "):
            add_runs(doc.add_heading("", 2), stripped[4:])
        elif stripped.startswith(">"):
            quoted = stripped.lstrip(">").strip()
            if not quoted:
                continue
            if quoted.startswith("### "):
                add_quote(doc, quoted[4:], heading=True)
            else:
                add_quote(doc, quoted)
        elif stripped.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        else:
            add_runs(doc.add_paragraph(), stripped)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
